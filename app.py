from flask import Flask, request, send_file, jsonify
import os
import json
import tempfile
import traceback
import shutil
import time
from io import BytesIO
from werkzeug.utils import secure_filename
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pdf2docx import Converter
from docx2pdf import convert
import subprocess
from xml.etree.ElementTree import Element, SubElement, tostring
import xml.dom.minidom
from PIL import Image  # Pillow kütüphanesinden
import PyPDF2  # PyPDF2 kütüphanesinden

# PAKET DÖNÜŞÜMLERİNDE KENDİ LOKALİMDE FARKLI LINUX DA FARKLI PAKET KULLANIYORUM DİKKAT ET

# Dosya Dönüştürme Servisi sınıfı
class FileConverterService:
    def __init__(self):
        self.supported_conversions = [
            'pdf_to_docx',
            'docx_to_pdf',
            'docx_to_json',
            'docx_to_excel',
            'json_to_excel',
            'json_to_docx',
            'json_to_pdf',
            'excel_to_json',
            'excel_to_docx',
            'excel_to_pdf',
            'csv_to_excel',
            'excel_to_xml',
            'png_to_pdf',  # Yeni eklenen dönüşüm türleri
            'jpg_to_pdf',
            'merge_pdfs'
        ]
    # PDF'den DOCX'e dönüşüm
    def pdf_to_docx(self, input_path, output_path):
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
        return output_path
    
    def docx_to_pdf(self, input_path, output_path):
        try:
            # Girdi dizini ve dosya adını al
            input_dir = os.path.dirname(input_path)
            input_filename = os.path.splitext(os.path.basename(input_path))[0]
            
            # LibreOffice ile dönüşüm
            result = subprocess.run([
                'libreoffice', 
                '--headless', 
                '--convert-to', 
                'pdf', 
                '--outdir', 
                input_dir,  # Aynı dizine kaydet
                input_path
            ], capture_output=True, text=True, check=True)
            
            # Olası PDF dosya yolları
            possible_pdf_paths = [
                os.path.join(input_dir, f"{input_filename}.pdf"),  # Linux tarzı
                os.path.join(input_dir, "input.pdf"),  # Alternatif
                os.path.join(input_dir, f"{os.path.basename(input_path)}.pdf")  # Başka bir alternatif
            ]
            
            # PDF dosyasını bul ve hedef konuma kopyala
            for pdf_path in possible_pdf_paths:
                if os.path.exists(pdf_path):
                    shutil.copy(pdf_path, output_path)
                    return output_path
            
            # Eğer hiçbir PDF bulunamazsa detaylı hata mesajı
            raise FileNotFoundError(f"PDF dosyası oluşturulamadı. LibreOffice çıktısı: {result.stdout}")
        
        except subprocess.CalledProcessError as e:
            raise Exception(f"LibreOffice dönüşüm hatası: {e.stderr}")
        except Exception as e:
            raise Exception(f"Dönüşüm sırasında hata oluştu: {str(e)}")
    
    # CSV'den Excel'e dönüşüm
    def csv_to_excel(self, input_path, output_path):
        try:
            # CSV dosyasını oku (encoding ve ayırıcı karakteri otomatik tespit etmeye çalış)
            try:
                # İlk olarak UTF-8 encoding ile dene
                df = pd.read_csv(input_path, encoding='utf-8')
            except UnicodeDecodeError:
                # UTF-8 çalışmazsa ISO-8859-1 encoding ile dene
                df = pd.read_csv(input_path, encoding='ISO-8859-1')
            except Exception:
                # Farklı ayırıcı karakteri olan CSV dosyalarını tespit et
                try:
                    df = pd.read_csv(input_path, encoding='utf-8', sep=';')
                except:
                    df = pd.read_csv(input_path, encoding='ISO-8859-1', sep=';')
            
            # Excel'e kaydet (varsayılan olarak ilk sayfa adı "CSV Data" olacak)
            with pd.ExcelWriter(output_path) as writer:
                df.to_excel(writer, sheet_name='CSV Data', index=False)
            
            return output_path
        
        except Exception as e:
            raise Exception(f"CSV'den Excel'e dönüşüm sırasında hata oluştu: {str(e)}")
    
    # DOCX'den JSON'a dönüşüm
    def docx_to_json(self, input_path, output_path):
        # Word belgesini UTF-8 encoding ile oku
        doc = Document(input_path)
        
        # JSON için veri yapısı oluştur
        data = {
            "paragraphs": [p.text for p in doc.paragraphs],
            "tables": []
        }
        
        # Tablo verilerini al
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_rows.append(row_data)
            data["tables"].append(table_rows)
        
        # JSON olarak UTF-8 ile kaydet
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)  # Türkçe karakterleri korumak için
        
        return output_path
    
    # DOCX'den Excel'e dönüşüm
    def docx_to_excel(self, input_path, output_path):
        # Word belgesini oku
        doc = Document(input_path)
        
        # Paragraf dataframe'i oluştur
        paragraphs_df = pd.DataFrame({"Content": [p.text for p in doc.paragraphs if p.text.strip()]})
        
        # Excel'e kaydet
        with pd.ExcelWriter(output_path) as writer:
            # Paragrafları kaydet
            if not paragraphs_df.empty:
                paragraphs_df.to_excel(writer, sheet_name='Paragraphs', index=False)
            
            # Tabloları kaydet
            for i, table in enumerate(doc.tables):
                # Maksimum sütun sayısını al
                max_cols = max(len(row.cells) for row in table.rows)
                
                # Başlıkları oluştur
                headers = [f"Column_{j+1}" for j in range(max_cols)]
                
                # Tablo verilerini al
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    # Gerekirse boş hücreler ekle
                    row_data.extend([''] * (max_cols - len(row_data)))
                    table_data.append(row_data)
                
                # Dataframe oluştur ve Excel'e kaydet
                if table_data:
                    df = pd.DataFrame(table_data, columns=headers)
                    sheet_name = f'Table_{i+1}'[:31].replace('/', '_').replace('\\', '_')
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        return output_path
        
    # JSON'dan Excel'e dönüşüm
    def json_to_excel(self, input_path, output_path):
        # JSON dosyasını oku
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Excel'e kaydet
        with pd.ExcelWriter(output_path) as writer:
            if isinstance(data, list):
                # Nesne listesini işle
                df = pd.DataFrame(data)
                df.to_excel(writer, sheet_name='Data', index=False)
            elif isinstance(data, dict):
                # Sözlüğü işle
                for key, value in data.items():
                    if isinstance(value, list):
                        df = pd.DataFrame(value)
                    else:
                        df = pd.DataFrame([value])
                    
                    # Sayfa adını düzelt
                    sheet_name = str(key)[:31].replace('/', '_').replace('\\', '_').replace('?', '_').replace('*', '_')
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        return output_path
    
    # JSON'dan DOCX'e dönüşüm
    def json_to_docx(self, input_path, output_path):
        # JSON dosyasını UTF-8 encoding ile oku
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Word belgesi oluştur
        doc = Document()
        doc.add_heading('JSON Verisi', 0)
        
        # JSON verilerini işle
        self._add_json_to_docx(doc, data)
        
        # Belgeyi UTF-8 ile kaydet
        doc.save(output_path)
        return output_path
    
    def _add_json_to_docx(self, doc, data, level=1):
        # Sözlükleri işle
        if isinstance(data, dict):
            for key, value in data.items():
                doc.add_heading(str(key), level)
                if isinstance(value, (dict, list)):
                    self._add_json_to_docx(doc, value, level + 1)
                else:
                    doc.add_paragraph(str(value))
        # Listeleri işle
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    doc.add_heading(f"Item {i+1}", level)
                    self._add_json_to_docx(doc, item, level + 1)
                else:
                    doc.add_paragraph(str(item))
        # Basit değerleri işle
        else:
            doc.add_paragraph(str(data))
    
    # JSON'dan PDF'e dönüşüm
    def json_to_pdf(self, input_path, output_path):
        # JSON dosyasını oku
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # PDF oluştur
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Başlık ekle
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "JSON Data")
        
        # JSON'ı string olarak formatla
        json_str = json.dumps(data, indent=2, ensure_ascii=False)
        lines = json_str.split('\n')
        
        # Her satırı PDF'e ekle
        y_pos = height - 80
        c.setFont("Helvetica", 12)
        
        for line in lines:
            # Gerekirse yeni sayfa oluştur
            if y_pos <= 50:
                c.showPage()
                y_pos = height - 50
            
            c.drawString(50, y_pos, line)
            y_pos -= 15
        
        c.save()
        return output_path
    
    # Excel'den JSON'a dönüşüm
    def excel_to_json(self, input_path, output_path):
        # Excel dosyasını UTF-8 ile okuma
        excel = pd.ExcelFile(input_path, encoding='utf-8')
        
        # Sonuç sözlüğü oluştur
        result = {}
        
        # Her sayfayı işle
        for sheet in excel.sheet_names:
            df = pd.read_excel(excel, sheet)
            # Türkçe karakterleri koruyarak JSON'a dönüştür
            result[sheet] = json.loads(df.to_json(orient='records', date_format='iso', force_ascii=False))
        
        # JSON dosyasına UTF-8 ile kaydet
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=4, ensure_ascii=False)
    
    # Excel'den DOCX'e dönüşüm
    def excel_to_docx(self, input_path, output_path):
        # Excel dosyasını oku
        excel = pd.ExcelFile(input_path)
        
        # Word belgesi oluştur
        doc = Document()
        doc.add_heading('Excel Data', 0)
        
        # Her sayfayı işle
        for sheet in excel.sheet_names:
            df = pd.read_excel(excel, sheet)
            
            # Sayfa başlığını ekle
            doc.add_heading(sheet, level=1)
            
            # Tablo oluştur
            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Başlıkları ekle
            for j, col in enumerate(df.columns):
                table.cell(0, j).text = str(col)
            
            # Verileri ekle
            for i, row in enumerate(df.itertuples()):
                for j in range(len(df.columns)):
                    table.cell(i + 1, j).text = str(row[j + 1])
            
            # Tablodan sonra boşluk ekle
            doc.add_paragraph()
        
        # Belgeyi kaydet
        doc.save(output_path)
        return output_path
    
    # Excel'den PDF'e dönüşüm
    def excel_to_pdf(self, input_path, output_path):
        # Excel dosyasını oku
        excel = pd.ExcelFile(input_path)
        
        # PDF oluştur
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Her sayfayı işle
        for sheet in excel.sheet_names:
            df = pd.read_excel(excel, sheet)
            
            # Sayfa başlığını ekle
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, sheet)
            
            # Parametreleri ayarla
            y_pos = height - 80
            row_height = 20
            col_widths = []
            
            # Sütun genişliklerini hesapla
            # Sütun genişliklerini hesapla
            for col in df.columns:
                max_width = max(
                    len(str(col)) * 7,
                    df[col].astype(str).map(len).max() * 7
                )
                col_widths.append(min(max_width, 150))
            
            # Başlık satırını ekle
            c.setFont("Helvetica-Bold", 12)
            x_pos = 50
            for i, col in enumerate(df.columns):
                c.drawString(x_pos, y_pos, str(col))
                x_pos += col_widths[i]
            
            # Veri satırlarını ekle
            y_pos -= row_height
            c.setFont("Helvetica", 10)
            
            for _, row in df.iterrows():
                # Gerekirse yeni sayfa oluştur
                if y_pos <= 50:
                    c.showPage()
                    y_pos = height - 50
                    
                    # Başlıkları yeniden çiz
                    c.setFont("Helvetica-Bold", 12)
                    x_pos = 50
                    for i, col in enumerate(df.columns):
                        c.drawString(x_pos, y_pos, str(col))
                        x_pos += col_widths[i]
                    
                    y_pos -= row_height
                    c.setFont("Helvetica", 10)
                
                # Satır verilerini çiz
                x_pos = 50
                for i, col in enumerate(df.columns):
                    c.drawString(x_pos, y_pos, str(row[col]))
                    x_pos += col_widths[i]
                
                y_pos -= row_height
            
            # Sonraki sayfa için yeni sayfa ekle
            c.showPage()
        
        c.save()
        return output_path
    
        # Excel'den XML'e dönüşüm

    # Excel'den XML'e dönüşüm
    def excel_to_xml(self, input_path, output_path):
        try:
            # Excel dosyasını oku
            df = pd.read_excel(input_path)
            
            # Kök XML elementini oluştur
            root = Element('root')
            
            # Her satırı XML'e dönüştür
            for _, row in df.iterrows():
                row_element = SubElement(root, 'row')
                
                # Her kolonu XML alt elementi olarak ekle
                for col_name, value in row.items():
                    # NaN değerlerini atla
                    if pd.isna(value):
                        continue
                        
                    # Kolon adını geçerli XML etiketi haline getir (boşluk ve özel karakterleri değiştir)
                    col_name = ''.join(c if c.isalnum() else '_' for c in str(col_name))
                    if col_name[0].isdigit():
                        col_name = 'f_' + col_name
                        
                    # Elementi oluştur ve metin değerini ayarla
                    col_element = SubElement(row_element, col_name)
                    col_element.text = str(value)
            
            # XML dizesini güzel formatlı olarak oluştur
            rough_string = tostring(root, 'utf-8')
            reparsed = xml.dom.minidom.parseString(rough_string)
            pretty_xml = reparsed.toprettyxml(indent="  ")
            
            # Dosyaya yaz
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(pretty_xml)
                
            return output_path
            
        except Exception as e:
            raise Exception(f"Excel'den XML'e dönüşüm sırasında hata oluştu: {str(e)}")

    # PNG'den PDF'e dönüşüm
    def png_to_pdf(self, input_path, output_path):
        try:
            # Resmi aç
            image = Image.open(input_path)
            
            # PDF oluştur
            c = canvas.Canvas(output_path, pagesize=letter)
            
            # Resmin boyutlarını al
            img_width, img_height = image.size
            
            # Resmi PDF sayfasına sığdır
            page_width, page_height = letter
            width_ratio = page_width / img_width
            height_ratio = page_height / img_height
            scale_ratio = min(width_ratio, height_ratio) * 0.9  # %90 oranında küçült
            
            new_width = img_width * scale_ratio
            new_height = img_height * scale_ratio
            
            # Resmi ortalamak için koordinatları hesapla
            x_centered = (page_width - new_width) / 2
            y_centered = (page_height - new_height) / 2
            
            # Resmi PDF'e ekle
            c.drawImage(input_path, x_centered, y_centered, width=new_width, height=new_height)
            
            c.save()
            return output_path
        except Exception as e:
            raise Exception(f"PNG'den PDF'e dönüşüm sırasında hata oluştu: {str(e)}")
   
    # PDF dosyalarını birleştirme
    def merge_pdfs(self, input_paths, output_path):
        try:
            # PDF birleştirici oluştur
            merger = PyPDF2.PdfMerger()
            
            # Benzersiz dosya içeriklerini kontrol etmek için
            seen_hashes = set()
            
            for pdf_path in input_paths:
                if not os.path.exists(pdf_path):
                    raise FileNotFoundError(f"PDF dosyası bulunamadı: {pdf_path}")
                
                # Dosya içeriğinin hash'ini al
                with open(pdf_path, 'rb') as f:
                    file_hash = hash(f.read())
                    
                    # Eğer bu hash daha önce görülmüşse atla
                    if file_hash in seen_hashes:
                        print(f"Uyarı: Yinelenen içerik atlandı: {pdf_path}")
                        continue
                        
                    seen_hashes.add(file_hash)
                    f.seek(0)  # Dosya pozisyonunu sıfırla
                    merger.append(f)
            
            # Birleştirilmiş PDF'i kaydet
            with open(output_path, 'wb') as f:
                merger.write(f)
            
            return output_path
        except Exception as e:
            raise Exception(f"PDF birleştirme sırasında hata oluştu: {str(e)}")
        finally:
            if 'merger' in locals():
                merger.close()
    
    # JPG'den PDF'e dönüşüm
    def jpg_to_pdf(self, input_path, output_path):
        try:
            # Resmi aç
            image = Image.open(input_path)
            
            # PDF oluştur
            c = canvas.Canvas(output_path, pagesize=letter)
            
            # Resmin boyutlarını al
            img_width, img_height = image.size
            
            # Resmi PDF sayfasına sığdır
            page_width, page_height = letter
            width_ratio = page_width / img_width
            height_ratio = page_height / img_height
            scale_ratio = min(width_ratio, height_ratio) * 0.9  # %90 oranında küçült
            
            new_width = img_width * scale_ratio
            new_height = img_height * scale_ratio
            
            # Resmi ortalamak için koordinatları hesapla
            x_centered = (page_width - new_width) / 2
            y_centered = (page_height - new_height) / 2
            
            # Resmi PDF'e ekle
            c.drawImage(input_path, x_centered, y_centered, width=new_width, height=new_height)
            
            c.save()
            return output_path
        except Exception as e:
            raise Exception(f"JPG'den PDF'e dönüşüm sırasında hata oluştu: {str(e)}")


# Flask uygulaması oluştur
app = Flask(__name__, static_folder='static', static_url_path='')

# Dönüştürücü servisi oluştur
converter_service = FileConverterService()

# Ana sayfa yolu
@app.route('/')
def index():
    return app.send_static_file('index.html')

# Dosya dönüştürme API uç noktası
@app.route('/api/convert', methods=['POST'])
def convert_file():
    temp_dir = None
    try:
        if request.form.get('conversionType') == 'merge_pdfs':
            if 'files' not in request.files:
                return jsonify({'error': 'PDF dosyaları bulunamadı'}), 400
            
            files = request.files.getlist('files')
            if len(files) < 2:
                return jsonify({'error': 'En az 2 PDF dosyası gereklidir'}), 400
            
            # Geçici dizin oluştur
            temp_dir = tempfile.mkdtemp()
            input_paths = []
            
            # Dosyaları kaydet ve benzersiz isimler oluştur
            for i, file in enumerate(files):
                if not file.filename.lower().endswith('.pdf'):
                    return jsonify({'error': 'Sadece PDF dosyaları yükleyebilirsiniz'}), 400
                
                # Dosya adını ve zaman damgasını kullanarak benzersiz isim oluştur
                timestamp = int(time.time() * 1000)
                unique_name = f"{timestamp}_{i}_{secure_filename(file.filename)}"
                input_path = os.path.join(temp_dir, unique_name)
                file.save(input_path)
                input_paths.append(input_path)
            
            # Çıkış dosyası yolu
            output_path = os.path.join(temp_dir, "merged.pdf")
            
            # Birleştirme metodunu çağır
            converter_service.merge_pdfs(input_paths, output_path)
            
            # Dosyayı gönder
            return send_file(
                output_path,
                as_attachment=True,
                download_name="merged.pdf",
                mimetype='application/pdf'
            )

        # Diğer dönüşüm türleri için mevcut kod
        if 'file' not in request.files:
            return jsonify({'error': 'Dosya bulunamadı'}), 400
        
        file = request.files['file']
        conversion_type = request.form.get('conversionType')
        
        if file.filename == '':
            return jsonify({'error': 'Dosya seçilmedi'}), 400
        
        if conversion_type not in converter_service.supported_conversions:
            return jsonify({'error': 'Desteklenmeyen dönüşüm türü'}), 400
        
        input_ext = file.filename.split('.')[-1].lower()
        output_ext = conversion_type.split('_to_')[1]
        
        ext_map = {
            'excel': 'xlsx',
            'docx': 'docx',
            'pdf': 'pdf',
            'json': 'json',
            'xml': 'xml',
            'png': 'png',
            'jpg': 'jpg'
        }
        if output_ext in ext_map:
            output_ext = ext_map[output_ext]
        
        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, f"input.{input_ext}")
        output_path = os.path.join(temp_dir, f"output.{output_ext}")
        
        file.save(input_path)
        
        convert_method = getattr(converter_service, conversion_type)
        convert_method(input_path, output_path)
        
        with open(output_path, 'rb') as f:
            file_data = f.read()
        
        file_stream = BytesIO(file_data)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{os.path.splitext(secure_filename(file.filename))[0]}.{output_ext}",
            mimetype='application/octet-stream'
        )
            
    except Exception as e:
        print(traceback.format_exc())  
        return jsonify({'error': f'Dönüşüm sırasında bir hata oluştu: {str(e)}'}), 500
    finally:
        if temp_dir and os.path.exists(temp_dir):
            try:
                time.sleep(0.5)
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                print(f"Temizleme hatası: {str(e)}")

@app.route('/kontrol')
def create_static_folder():
    try:
        if not os.path.exists('static'):
            os.mkdir('static')
        return "Static klasörü oluşturuldu", 200
    except Exception as e:
        return f"Hata: {str(e)}", 500

if __name__ == '__main__':
    if not os.path.exists('static'):
        os.mkdir('static')
    app.run(debug=True, host='0.0.0.0', port=5000)